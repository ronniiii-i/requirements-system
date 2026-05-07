import io
import json
import logging
from datetime import datetime
from uuid import UUID
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from app.models.requirement import Requirement
from app.models.user_project import Project
from app.models.enums import RequirementStatus

logger = logging.getLogger(__name__)


# ── Helpers ────────────────────────────────────────────────────────────────

def _get_requirements(project_id: UUID, db: Session) -> list[Requirement]:
    return (
        db.query(Requirement)
        .filter(
            Requirement.project_id == project_id,
            Requirement.is_current_version == True,
        )
        .order_by(Requirement.created_at.asc())
        .all()
    )


def _status_label(status: RequirementStatus) -> str:
    return status.value.replace("_", " ").title()


def _priority_label(priority) -> str:
    labels = {
        "must_have": "Must Have",
        "should_have": "Should Have",
        "could_have": "Could Have",
        "wont_have": "Won't Have",
    }
    return labels.get(priority.value, priority.value)


def _score_label(score: float | None) -> str:
    if score is None:
        return "N/A"
    pct = round(score * 100, 1)
    if pct >= 80:
        return f"{pct}% ✓"
    elif pct >= 60:
        return f"{pct}% ⚠"
    else:
        return f"{pct}% ✗"


# ── Word Export (IEEE 29148) ───────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_field_row(table, label: str, value: str, shade_label=True):
    """Add a two-column row to a requirement detail table."""
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]

    label_cell.text = label
    label_cell.paragraphs[0].runs[0].bold = True
    label_cell.paragraphs[0].runs[0].font.size = Pt(9)

    value_cell.text = value or "—"
    value_cell.paragraphs[0].runs[0].font.size = Pt(9)

    if shade_label:
        _set_cell_bg(label_cell, "EBF3FB")


def export_word(project: Project, db: Session) -> bytes:
    """
    Generates an IEEE 29148 compliant requirements document as .docx bytes.
    """
    requirements = _get_requirements(project.id, db)
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Cover page ────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("SOFTWARE REQUIREMENTS SPECIFICATION")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    project_para = doc.add_paragraph()
    project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project_run = project_para.add_run(project.name.upper())
    project_run.bold = True
    project_run.font.size = Pt(14)

    if project.domain:
        domain_para = doc.add_paragraph()
        domain_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        domain_para.add_run(f"Domain: {project.domain}").font.size = Pt(11)

    doc.add_paragraph()

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y')}\n"
        f"Standard: IEEE 29148:2018\n"
        f"Total Requirements: {len(requirements)}"
    ).font.size = Pt(10)

    doc.add_page_break()

    # ── Section 1: Introduction ───────────────────────────────────────────
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        f"This Software Requirements Specification (SRS) document describes the "
        f"functional and non-functional requirements for the {project.name} system. "
        f"This document was generated automatically by the ReqGen AI-powered requirements "
        f"generation system and is intended for review by requirement engineers and domain experts."
    )

    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph(
        f"This document covers {len(requirements)} requirements for the {project.name} system"
        + (f" in the {project.domain} domain." if project.domain else ".")
    )
    if project.description:
        doc.add_paragraph(project.description)

    doc.add_heading("1.3 Document Conventions", level=2)
    doc.add_paragraph(
        "This document follows IEEE 29148:2018 conventions. "
        "The keyword 'shall' indicates a mandatory requirement. "
        "Each requirement includes a unique identifier, statement, rationale, "
        "and fit criterion (acceptance test)."
    )

    doc.add_page_break()

    # ── Section 2: Requirements Summary ───────────────────────────────────
    doc.add_heading("2. Requirements Summary", level=1)

    summary_table = doc.add_table(rows=1, cols=5)
    summary_table.style = "Table Grid"

    headers = ["Req ID", "Title", "Type", "Priority", "Status"]
    header_row = summary_table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(cell, "1F497D")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for req in requirements:
        row = summary_table.add_row()
        row.cells[0].text = req.req_id
        row.cells[1].text = req.title[:60] + ("..." if len(req.title) > 60 else "")
        row.cells[2].text = req.type.value.replace("_", " ").title()
        row.cells[3].text = _priority_label(req.priority)
        row.cells[4].text = _status_label(req.status)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ── Section 3: Detailed Requirements ─────────────────────────────────
    doc.add_heading("3. Detailed Requirements", level=1)

    for req in requirements:
        # Requirement heading
        req_heading = doc.add_heading(f"{req.req_id}: {req.title}", level=2)

        # Detail table
        detail_table = doc.add_table(rows=0, cols=2)
        detail_table.style = "Table Grid"
        detail_table.columns[0].width = Cm(4)
        detail_table.columns[1].width = Cm(11.5)

        _add_field_row(detail_table, "Identifier", req.req_id)
        _add_field_row(detail_table, "Type", req.type.value.replace("_", " ").title())
        _add_field_row(detail_table, "Status", _status_label(req.status))
        _add_field_row(detail_table, "Priority", _priority_label(req.priority))
        _add_field_row(detail_table, "Statement", req.statement)
        _add_field_row(detail_table, "Rationale", req.rationale or "—")
        _add_field_row(detail_table, "Fit Criterion", req.fit_criterion or "—")
        _add_field_row(detail_table, "Originator", req.originator or "System (AI Generated)")
        _add_field_row(detail_table, "Version", str(req.version))

        # QA scores
        _add_field_row(
            detail_table,
            "Quality Score",
            f"Overall: {_score_label(req.overall_quality_score)}  |  "
            f"Ambiguity: {_score_label(req.ambiguity_score)}  |  "
            f"Completeness: {_score_label(req.completeness_score)}  |  "
            f"Testability: {_score_label(req.testability_score)}"
        )

        # QA issues
        if req.qa_issues:
            issues_text = "\n".join(
                f"[{i['severity'].upper()}] {i['message']}"
                for i in req.qa_issues
            )
            _add_field_row(detail_table, "QA Issues", issues_text)

        doc.add_paragraph()  # spacing between requirements

    doc.add_page_break()

    # ── Section 4: Traceability Matrix ────────────────────────────────────
    doc.add_heading("4. Requirements Traceability Matrix", level=1)

    rtm_table = doc.add_table(rows=1, cols=4)
    rtm_table.style = "Table Grid"

    rtm_headers = ["Req ID", "Statement (Summary)", "Type", "Quality Score"]
    rtm_header_row = rtm_table.rows[0]
    for i, h in enumerate(rtm_headers):
        cell = rtm_header_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(cell, "1F497D")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for req in requirements:
        row = rtm_table.add_row()
        row.cells[0].text = req.req_id
        stmt = req.statement[:80] + ("..." if len(req.statement) > 80 else "")
        row.cells[1].text = stmt
        row.cells[2].text = req.type.value.replace("_", " ").title()
        row.cells[3].text = (
            f"{round(req.overall_quality_score * 100)}%"
            if req.overall_quality_score else "N/A"
        )
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ── Excel Export ──────────────────────────────────────────────────────────

def export_excel(project: Project, db: Session) -> bytes:
    requirements = _get_requirements(project.id, db)

    wb = Workbook()

    # ── Sheet 1: Requirements ──────────────────────────────────────────
    ws = wb.active
    ws.title = "Requirements"

    header_font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
    header_fill = PatternFill("solid", fgColor="1F497D")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_align = Alignment(vertical="top", wrap_text=True)
    border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    alt_fill = PatternFill("solid", fgColor="EBF3FB")

    headers = [
        "Req ID", "Title", "Type", "Status", "Priority",
        "Statement", "Rationale", "Fit Criterion",
        "AI Confidence", "Ambiguity", "Completeness",
        "Consistency", "Testability", "Overall Quality",
        "Business Value", "Risk", "Cost/Effort",
        "Stakeholder Importance", "Weighted Score",
        "Version", "Created At",
    ]

    col_widths = [10, 30, 15, 14, 14, 50, 40, 40, 13, 12, 14, 13, 13, 15, 14, 8, 12, 22, 15, 9, 20]

    for col_idx, (header, width) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = border
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    ws.row_dimensions[1].height = 30

    for row_idx, req in enumerate(requirements, 2):
        fill = alt_fill if row_idx % 2 == 0 else PatternFill()
        values = [
            req.req_id,
            req.title,
            req.type.value.replace("_", " ").title(),
            _status_label(req.status),
            _priority_label(req.priority),
            req.statement,
            req.rationale or "",
            req.fit_criterion or "",
            round(req.ai_confidence * 100, 1) if req.ai_confidence else None,
            round(req.ambiguity_score * 100, 1) if req.ambiguity_score else None,
            round(req.completeness_score * 100, 1) if req.completeness_score else None,
            round(req.consistency_score * 100, 1) if req.consistency_score else None,
            round(req.testability_score * 100, 1) if req.testability_score else None,
            round(req.overall_quality_score * 100, 1) if req.overall_quality_score else None,
            req.business_value_score,
            req.risk_score,
            req.cost_effort_score,
            req.stakeholder_importance,
            float(req.weighted_score) if req.weighted_score else None,
            req.version,
            req.created_at.strftime("%Y-%m-%d %H:%M") if req.created_at else "",
        ]

        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = cell_align
            cell.border = border
            if fill.fgColor.rgb != "00000000":
                cell.fill = fill

        ws.row_dimensions[row_idx].height = 40

    ws.freeze_panes = "A2"

    # ── Sheet 2: QA Issues ─────────────────────────────────────────────
    ws2 = wb.create_sheet("QA Issues")
    issue_headers = ["Req ID", "Issue Type", "Severity", "Message", "Suggestion"]
    issue_widths = [10, 15, 10, 50, 50]

    for col_idx, (header, width) in enumerate(zip(issue_headers, issue_widths), 1):
        cell = ws2.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = border
        ws2.column_dimensions[get_column_letter(col_idx)].width = width

    issue_row = 2
    severity_colors = {"high": "FFD7D7", "medium": "FFF3CD", "low": "D4EDDA"}

    for req in requirements:
        if req.qa_issues:
            for issue in req.qa_issues:
                sev = issue.get("severity", "low")
                row_fill = PatternFill("solid", fgColor=severity_colors.get(sev, "FFFFFF"))
                row_values = [
                    req.req_id,
                    issue.get("type", ""),
                    sev.upper(),
                    issue.get("message", ""),
                    issue.get("suggestion", ""),
                ]
                for col_idx, value in enumerate(row_values, 1):
                    cell = ws2.cell(row=issue_row, column=col_idx, value=value)
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    cell.border = border
                    cell.fill = row_fill
                issue_row += 1

    ws2.freeze_panes = "A2"

    # ── Sheet 3: Summary stats ─────────────────────────────────────────
    ws3 = wb.create_sheet("Summary")
    ws3.column_dimensions["A"].width = 30
    ws3.column_dimensions["B"].width = 20

    summary_title = ws3.cell(row=1, column=1, value=f"Requirements Summary — {project.name}")
    summary_title.font = Font(bold=True, size=14, color="1F497D", name="Arial")

    stats = [
        ("Total Requirements", len(requirements)),
        ("Approved", sum(1 for r in requirements if r.status.value == "approved")),
        ("Under Review", sum(1 for r in requirements if r.status.value == "under_review")),
        ("Draft", sum(1 for r in requirements if r.status.value == "draft")),
        ("Rejected", sum(1 for r in requirements if r.status.value == "rejected")),
        ("", ""),
        ("Functional", sum(1 for r in requirements if r.type.value == "functional")),
        ("Non-Functional", sum(1 for r in requirements if r.type.value == "non_functional")),
        ("Performance", sum(1 for r in requirements if r.type.value == "performance")),
        ("Security", sum(1 for r in requirements if r.type.value == "security")),
        ("", ""),
        ("Avg Quality Score", f"{round(sum(r.overall_quality_score or 0 for r in requirements) / max(len(requirements), 1) * 100, 1)}%"),
        ("Avg AI Confidence", f"{round(sum(r.ai_confidence or 0 for r in requirements) / max(len(requirements), 1) * 100, 1)}%"),
        ("Requirements with QA Issues", sum(1 for r in requirements if r.qa_issues)),
    ]

    for row_idx, (label, value) in enumerate(stats, 3):
        if label:
            label_cell = ws3.cell(row=row_idx, column=1, value=label)
            label_cell.font = Font(bold=True, name="Arial", size=10)
            ws3.cell(row=row_idx, column=2, value=value)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ── JSON Export ───────────────────────────────────────────────────────────

def export_json(project: Project, db: Session) -> dict:
    requirements = _get_requirements(project.id, db)

    return {
        "export_metadata": {
            "generated_at": datetime.now().isoformat(),
            "standard": "IEEE 29148:2018",
            "system": "ReqGen",
            "project": project.name,
            "domain": project.domain,
            "total_requirements": len(requirements),
        },
        "requirements": [
            {
                "req_id": req.req_id,
                "title": req.title,
                "statement": req.statement,
                "type": req.type.value,
                "status": req.status.value,
                "priority": req.priority.value,
                "rationale": req.rationale,
                "fit_criterion": req.fit_criterion,
                "originator": req.originator,
                "version": req.version,
                "ai_generated": req.ai_generated,
                "ai_confidence": req.ai_confidence,
                "quality_scores": {
                    "overall": req.overall_quality_score,
                    "ambiguity": req.ambiguity_score,
                    "completeness": req.completeness_score,
                    "consistency": req.consistency_score,
                    "testability": req.testability_score,
                },
                "qa_issues": req.qa_issues or [],
                "prioritization": {
                    "business_value": req.business_value_score,
                    "risk": req.risk_score,
                    "cost_effort": req.cost_effort_score,
                    "stakeholder_importance": req.stakeholder_importance,
                    "weighted_score": float(req.weighted_score) if req.weighted_score else None,
                },
                "created_at": req.created_at.isoformat() if req.created_at else None,
                "updated_at": req.updated_at.isoformat() if req.updated_at else None,
            }
            for req in requirements
        ],
    }